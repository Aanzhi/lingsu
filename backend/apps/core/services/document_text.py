"""Extract readable text from student-uploaded material attachments.

Supported sources:
  * plain text / markdown (.txt, .md)
  * Word documents (.docx) via python-docx
  * PDF (.pdf) via pypdf
  * images (.png/.jpg/.jpeg/.bmp/.tiff/.gif) via Pillow + pytesseract OCR

The Celery upload-scan task calls :func:`extract_attachment_text` after the
malware scan finishes. Extraction failures are isolated: a missing optional
dependency or an unreadable file degrades to ``UNSUPPORTED``/``FAILED`` instead
of crashing the upload pipeline.
"""
from __future__ import annotations

import io

from apps.core.models import MaterialAttachment

# Upper bound on bytes we attempt to feed an extractor. Keeps memory and OCR
# latency sane for very large uploads; downstream context injection truncates
# further.
_MAX_BYTES = 20 * 1024 * 1024

TEXT_EXTS = {".txt", ".md", ".text", ".csv", ".json", ".log"}
DOCX_EXTS = {".docx"}
PDF_EXTS = {".pdf"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".bmp", ".tif", ".tiff", ".gif", ".webp"}


def _ext(original_name: str) -> str:
    name = (original_name or "").lower()
    return "." + name.rsplit(".", 1)[-1] if "." in name else ""


def _read_bytes(file_field) -> bytes:
    with file_field.open("rb") as fh:
        return fh.read(_MAX_BYTES)


def _extract_text_bytes(data: bytes, ext: str) -> tuple[str, str | None]:
    """Return (text, error_or_none). error None means success."""
    if ext in TEXT_EXTS:
        try:
            return data.decode("utf-8-sig", errors="replace"), None
        except Exception as exc:  # pragma: no cover - defensive
            return "", f"文本读取失败：{exc}"

    if ext in DOCX_EXTS:
        try:
            from docx import Document
        except ImportError:
            return "", "未安装 python-docx，无法解析 Word 文档"
        try:
            document = Document(io.BytesIO(data))
            parts = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    cells = [c.text for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts), None
        except Exception as exc:
            return "", f"Word 解析失败：{exc}"

    if ext in PDF_EXTS:
        try:
            from pypdf import PdfReader
        except ImportError:
            return "", "未安装 pypdf，无法解析 PDF 文档"
        try:
            reader = PdfReader(io.BytesIO(data))
            pages = [p.extract_text() or "" for p in reader.pages]
            return "\n".join(pages), None
        except Exception as exc:
            return "", f"PDF 解析失败：{exc}"

    if ext in IMAGE_EXTS:
        try:
            from PIL import Image
            import pytesseract
        except ImportError:
            return "", "未安装 Pillow/pytesseract，无法识别图片文字"
        try:
            image = Image.open(io.BytesIO(data))
            # 优先中文，回退英文；任一语言引擎缺失时不致命
            text = pytesseract.image_to_string(image, lang="chi_sim+eng")
            return text, None
        except Exception as exc:
            return "", f"图片 OCR 失败：{exc}"

    return "", "暂不支持的文件类型"


def extract_attachment_text(file_field, original_name: str) -> tuple[str, str, str]:
    """Extract text for ``file_field``.

    Returns ``(text, extract_status_value, detail)`` where ``extract_status_value``
    is one of :class:`MaterialAttachment.ExtractStatus`.
    """
    ext = _ext(original_name)
    if ext not in (TEXT_EXTS | DOCX_EXTS | PDF_EXTS | IMAGE_EXTS):
        return "", MaterialAttachment.ExtractStatus.UNSUPPORTED, "暂不支持的文件类型"
    try:
        data = _read_bytes(file_field)
    except Exception as exc:
        return "", MaterialAttachment.ExtractStatus.FAILED, f"文件读取失败：{exc}"
    text, error = _extract_text_bytes(data, ext)
    if error:
        # Unsupported types are not failures — keep them distinct for the UI.
        if ext not in (TEXT_EXTS | DOCX_EXTS | PDF_EXTS | IMAGE_EXTS):
            return "", MaterialAttachment.ExtractStatus.UNSUPPORTED, error
        return "", MaterialAttachment.ExtractStatus.FAILED, error
    text = text.strip()
    if not text:
        return "", MaterialAttachment.ExtractStatus.UNSUPPORTED, "未抽取到可读文字（可能是扫描件或空白文件）"
    return text, MaterialAttachment.ExtractStatus.DONE, ""
