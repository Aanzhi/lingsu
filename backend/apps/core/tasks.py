from io import BytesIO
import hashlib
import json
import os
import tempfile
from datetime import timedelta

import requests
from openai import OpenAI
from celery import shared_task
from django.conf import settings
from django.core.files.base import ContentFile
from django.db import transaction
from django.utils import timezone
from docx import Document

from .models import AgentTemplate, AIConversationMessage
from .ai_agents import PAPER_AGENT_KEYS, parse_research_question_output, render_agent_prompt, render_conversation_agent_prompt
from .workflows.ai import publish_conversation_event

DEFAULT_AI_INSTRUCTION = (
    "你是青少年科创项目教练。只提供可编辑建议，不虚构数据、引用或实验结果，"
    "不替用户提交、审核或发布。明确指出需要学生核实的事实。"
)


def _ref_label(source: dict) -> str:
    kind = source.get("kind")
    title = source.get("title", "")
    if kind == "task":
        return f"步骤《{title}》"
    if kind == "attachment":
        return f"文件《{title}》"
    if kind == "ai_history":
        return f"既有 AI 草稿《{title}》"
    if kind == "teacher_feedback":
        return f"教师反馈《{title}》"
    return f"材料《{title}》"


def _research_question_artifact(output):
    structured = parse_research_question_output(output)
    if not structured:
        return None
    readable = "\n\n".join(
        f"候选 {index + 1}：{candidate['question']}\n研究边界：{candidate['scope']}\n为什么值得研究：{candidate['why']}\n证据或数据：{candidate['evidence_plan']}\n可能限制：{candidate['limitations']}"
        for index, candidate in enumerate(structured["candidates"])
    )
    return {
        "artifact_payload": {
            "title": "研究问题候选",
            "draft": readable,
            "content": output,
            "project_title": structured["project_title"],
            "project_type": structured["project_type"],
            "project_plan": structured["project_plan"],
            "candidates": structured["candidates"],
            "recommended_index": structured["recommended_index"],
            "missing_information": structured["missing_information"],
            "next_action": "请比较候选并结合真实可获得的证据核验后，再由你确认保存。",
        },
        "verification_items": [
            {"item": "研究对象、边界与证据来源", "status": "needs_verification", "guidance": "确认每个候选都能在你的时间、设备和样本条件下验证。"},
            {"item": "事实、数据与限制", "status": "needs_verification", "guidance": "AI 不会替你编造数据；保存前请检查限制是否准确。"},
        ],
    }


def _artifact_fields(record, output):
    """Build an auditable, UI-friendly artifact without replacing the raw model output."""
    if record.agent_key == "proposal-topic":
        structured = _research_question_artifact(output)
        if structured:
            return structured
    title = record.material.title if record.material_id else (record.purpose or "AI 生成草稿")
    return {
        "artifact_payload": {
            "title": title,
            "draft": output,
            "content": output,
            "next_action": "请结合原始资料逐项核验文献、数据和事实后再编辑或提交。",
        },
        "verification_items": [
            {
                "item": "文献、数据与事实依据",
                "status": "needs_verification",
                "guidance": "AI 输出仅为可编辑草稿；请使用真实、可追溯的原始资料完成核验。",
            }
        ],
    }


def _demo_ai_response(record, context_parts, referenced=None, agent_prompt=""):
    """Honest, clearly-labeled illustrative reply used when no OpenAI key is set.

    It references the real project context so the demo proves the wiring works,
    but never fabricates data, citations, or experimental results. For the
    cross-consistency agent it returns a strict JSON array matching the contract
    the frontend ConsistencyCheckCard parses.
    """
    project = record.project
    if record.agent_key == "proposal-topic":
        seed = (record.prompt or (project.problem if project else "") or (project.title if project else "") or "你的研究现象").strip()
        payload = {
            "candidates": [
                {
                    "question": f"在明确的研究对象与场景中，{seed}的主要表现和可能原因是什么？",
                    "scope": "先限定一个具体对象、场景和观察时段",
                    "why": "能够把兴趣转成可观察的现象，适合作为第一轮研究切口。",
                    "evidence_plan": "现场观察、记录表、访谈或公开资料；具体选择需由学生核验。",
                    "limitations": "样本量和观察周期可能有限，不能直接推断更大范围。",
                    "scores": {"researchability": 4, "clarity": 4, "verifiability": 4, "resource_fit": 4},
                },
                {
                    "question": f"哪些可观察条件与{seed}的差异相关？",
                    "scope": "比较两个或多个可获得条件，保持其他条件尽量一致",
                    "why": "便于形成对照，帮助区分相关现象与可能影响因素。",
                    "evidence_plan": "对照观察、简单测量或小规模调查，并记录原始数据。",
                    "limitations": "对照条件难以完全一致，结果只能支持有限范围的判断。",
                    "scores": {"researchability": 4, "clarity": 3, "verifiability": 4, "resource_fit": 3},
                },
                {
                    "question": f"在现有时间和资源下，如何验证{seed}的一种可行解释？",
                    "scope": "只验证一个解释和一项可测指标，不扩展到完整解决方案",
                    "why": "直接连接可执行的验证行动，适合资源有限的学生项目。",
                    "evidence_plan": "制定一次小实验或连续观察，预先写明指标、步骤和停止条件。",
                    "limitations": "一次验证不能证明普遍规律，需要明确结果的适用边界。",
                    "scores": {"researchability": 5, "clarity": 4, "verifiability": 4, "resource_fit": 4},
                },
            ],
            "recommended_index": 0,
            "missing_information": ["还需要确认研究对象、观察时段和可用设备。"],
        }
        return json.dumps(payload, ensure_ascii=False, indent=2)
    if record.agent_key == "cross-consistency":
        materials = list(project.materials.all().order_by("report_order", "id"))
        if len(materials) < 2:
            return json.dumps({
                "coverage_score": 0,
                "missing_evidence": [],
                "conflicts": [],
                "issues": [],
            }, ensure_ascii=False)
        titles = [m.title for m in materials]
        demo_payload = {
            "coverage_score": 72,
            "missing_evidence": [
                "建议补充对照/重复实验数据，以提升结论可信度（示例提示，真实体检会基于你的材料判断）。",
            ],
            "conflicts": [],
            "issues": [
                {
                    "severity": "低",
                    "title": "示例条目：各材料的研究目标口径建议统一核对",
                    "involves": titles[:2],
                    "detail": "演示示例：通读材料时，建议确认开头确定的研究问题在后续材料中被一致复述，避免前后表述偏差。",
                    "suggestion": "在每一份材料开头复述项目研究问题，保持术语与目标口径一致。",
                }
            ],
        }
        return json.dumps(demo_payload, ensure_ascii=False, indent=2)
    if record.agent_key == "next-step-advisor":
        done = project.materials.filter(status="approved").count()
        total = project.materials.count()
        demo_payload = {
            "overall_progress": f"示例进度：已有 {done}/{total} 份材料通过，整体推进顺利，进入收尾与证据补齐阶段。",
            "next_actions": [
                {"priority": "高", "action": "把核心实验/观察的真实数据整理成规范记录表", "rationale": "数据是结论的支撑，越早定稿越稳", "related_task": "实验记录"},
                {"priority": "中", "action": "对照各步骤写作要求逐条核对是否遗漏评分要点", "rationale": "避免临时补写遗漏", "related_task": "研究报告起草"},
                {"priority": "低", "action": "准备一页答辩要点草稿", "rationale": "提前演练，降低临场压力", "related_task": "答辩问答准备"},
            ],
            "missing_evidence": [
                "建议补充对照/重复样本数据，提升结论可信度（示例提示，真实体检会基于你的材料判断）。",
            ],
            "risks": [
                "若引用外部文献或数据，请确认来源真实可查，AI 不会替你编造（示例提示）。",
            ],
        }
        return json.dumps(demo_payload, ensure_ascii=False, indent=2)
    lines = [
        "【演示模式 · 未接入真实大模型】",
        "以下为 AI 助手在读懂你的项目上下文后会给出的建议示例。配置 OPENAI_API_KEY 后即返回真实模型结果；"
        "演示内容不编造任何数据、引用或实验结果。",
        "",
        f"项目：{project.title}",
    ]
    if agent_prompt:
        lines.extend(["", "本次 Agent 输入（演示回显）：", agent_prompt])
    if record.task_id:
        try:
            t = record.task
            lines.append(f"当前步骤：{t.stage_name} · {t.title}")
        except Exception:
            pass
    if record.material_id:
        try:
            m = record.material
            rev = m.revisions.order_by("-created_at", "-id").first()
            if rev and rev.content:
                lines.append(f"你已写的内容（节选）：{rev.content[:400]}")
        except Exception:
            pass
    lines += [
        "",
        "建议（示例）：",
        "1. 先确认这一步要回答的核心问题，把已知事实列清楚再开始写作。",
        "2. 把你在实验或观察中得到的真实数据补充进去——AI 不会替你编造数据、引用或结论。",
        "3. 完成后对照本步骤的写作要求逐条核对，确保没有遗漏评分要点。",
        "4. 若有不确定的地方，把具体疑问写清楚再来问，AI 才能给更有针对性的建议。",
    ]
    if referenced:
        sources = "; ".join(_ref_label(r) for r in referenced[:8])
        lines.append("")
        lines.append(f"本次已读取的项目上下文来源：{sources}")
        files = [r for r in referenced if r.get("kind") == "attachment"]
        if files:
            lines.append(f"（其中已读取 {len(files)} 个上传文件的正文内容，用于辅助回答；真实模型接入后会基于这些原文作答。）")
    return "\n".join(lines)


def _finish_conversation_message(message, output, artifact=None):
    artifact = artifact or {"artifact_payload": {}, "verification_items": []}
    for chunk in [output[i:i + 240] for i in range(0, len(output), 240)]:
        publish_conversation_event(message.id, "message.delta", {"delta": chunk})
    message.content = output
    message.status = AIConversationMessage.Status.COMPLETED
    message.artifact_payload = artifact["artifact_payload"]
    message.save(update_fields=["content", "status", "artifact_payload", "updated_at"])
    publish_conversation_event(message.id, "message.artifact", artifact)
    publish_conversation_event(message.id, "message.done", {"message_id": message.id})


@shared_task(bind=True, autoretry_for=(Exception,), retry_backoff=True, max_retries=3)
def generate_general_ai_response(self, message_id):
    """Generate a real response for project-free conversations when configured."""
    message = AIConversationMessage.objects.select_related("conversation").get(pk=message_id)
    if message.status in {AIConversationMessage.Status.COMPLETED, AIConversationMessage.Status.FAILED}:
        return {"message_id": message.id, "status": message.status}
    message.status = AIConversationMessage.Status.STREAMING
    message.save(update_fields=["status", "updated_at"])
    publish_conversation_event(message.id, "message.started", {"message_id": message.id})
    try:
        conversation = message.conversation
        agent_key = conversation.current_agent or ""
        history = conversation.messages.exclude(pk=message.id).order_by("-created_at", "-id")[:8]
        transcript = "\n".join(
            f"{item.role}: {item.content}" for item in reversed(list(history)) if item.content
        )
        latest_user_message = next((item.content for item in history if item.role == AIConversationMessage.Role.USER and item.content), message.content)
        template = AgentTemplate.resolve(agent_key, conversation.owner.school, conversation.owner.role) if agent_key else None
        instructions = template.system_instruction if template else DEFAULT_AI_INSTRUCTION
        rendered_prompt = render_conversation_agent_prompt(template, conversation, latest_user_message) if template else latest_user_message
        if agent_key == "proposal-topic":
            instructions = (
                f"{instructions} 你正在执行研究问题助手工作流。"
                "只输出严格 JSON，不要 Markdown。JSON 必须包含 project_title、project_type、project_plan、"
                "candidates、recommended_index、missing_information。project_type 只能是 research、invention、engineering；"
                "candidates 必须正好 3 个，每个包含 question、scope、why、evidence_plan、limitations 和 scores，"
                "scores 的 researchability、clarity、verifiability、resource_fit 均为 1-5。"
            )
        client_kwargs = {"api_key": settings.OPENAI_API_KEY}
        if settings.OPENAI_BASE_URL:
            client_kwargs["base_url"] = settings.OPENAI_BASE_URL
        response = OpenAI(**client_kwargs).responses.create(
            model=settings.OPENAI_MODEL,
            instructions=instructions,
            input=(f"对话历史：\n{transcript}\n\n" if transcript else "") + f"Skill 指令：\n{rendered_prompt}\n用户问题：\n{latest_user_message}",
        )
        artifact = _research_question_artifact(response.output_text) if agent_key == "proposal-topic" else None
        _finish_conversation_message(message, response.output_text, artifact)
        return {"message_id": message.id, "status": message.status}
    except Exception as exc:
        message.status = AIConversationMessage.Status.FAILED
        message.error_message = str(exc)[:2000]
        message.save(update_fields=["status", "error_message", "updated_at"])
        publish_conversation_event(message.id, "message.error", {"error": message.error_message})
        raise


@shared_task(bind=True, autoretry_for=(Exception,), retry_backoff=True, max_retries=3)
def process_uploaded_material(self, revision_id):
    """Hash and malware-scan every attachment before it can be trusted."""
    from .models import MaterialAttachment, MaterialRevision

    revision = MaterialRevision.objects.prefetch_related("attachments").get(pk=revision_id)
    overall = MaterialAttachment.ScanStatus.CLEAN
    for attachment in revision.attachments.all():
        attachment.scan_status = MaterialAttachment.ScanStatus.PROCESSING
        attachment.scan_detail = ""
        attachment.save(update_fields=["scan_status", "scan_detail"])
        digest = hashlib.sha256()
        try:
            with attachment.file.open("rb") as source:
                for chunk in iter(lambda: source.read(1024 * 1024), b""):
                    digest.update(chunk)
            attachment.sha256 = digest.hexdigest()
            clamav_host = getattr(settings, "CLAMAV_HOST", "").strip()
            if clamav_host:
                import clamd

                scanner = clamd.ClamdNetworkSocket(
                    host=clamav_host,
                    port=getattr(settings, "CLAMAV_PORT", 3310),
                    timeout=getattr(settings, "CLAMAV_TIMEOUT", 120),
                )
                with attachment.file.open("rb") as source:
                    result = scanner.instream(source)
                verdict, detail = result.get("stream", ("ERROR", "扫描服务没有返回结果"))
                if verdict == "FOUND":
                    attachment.scan_status = MaterialAttachment.ScanStatus.INFECTED
                    attachment.scan_detail = str(detail)[:500]
                    overall = MaterialAttachment.ScanStatus.INFECTED
                elif verdict == "OK":
                    attachment.scan_status = MaterialAttachment.ScanStatus.CLEAN
                else:
                    raise RuntimeError(f"ClamAV 扫描失败：{detail}")
            elif getattr(settings, "FILE_SCAN_REQUIRED", False):
                raise RuntimeError("生产环境要求文件病毒扫描，但未配置 ClamAV。")
            else:
                attachment.scan_status = MaterialAttachment.ScanStatus.CLEAN
                attachment.scan_detail = "开发模式：已完成哈希校验，未启用病毒扫描。"
            attachment.scanned_at = timezone.now()
            attachment.save(update_fields=["sha256", "scan_status", "scan_detail", "scanned_at"])
            # 抽取附件可读文本（独立于病毒扫描，失败不应影响附件上传）
            try:
                from .services.document_text import extract_attachment_text

                text, ext_status, ext_detail = extract_attachment_text(attachment.file, attachment.original_name)
                attachment.extracted_text = text
                attachment.extract_status = ext_status
                attachment.extract_detail = ext_detail[:500]
                attachment.extracted_at = timezone.now()
                attachment.save(update_fields=["extracted_text", "extract_status", "extract_detail", "extracted_at"])
            except Exception as ext_exc:
                attachment.extract_status = MaterialAttachment.ExtractStatus.FAILED
                attachment.extract_detail = f"文本抽取异常：{ext_exc}"[:500]
                attachment.extracted_at = timezone.now()
                attachment.save(update_fields=["extract_status", "extract_detail", "extracted_at"])
        except Exception as exc:
            attachment.scan_status = MaterialAttachment.ScanStatus.FAILED
            attachment.scan_detail = str(exc)[:500]
            attachment.scanned_at = timezone.now()
            attachment.save(update_fields=["sha256", "scan_status", "scan_detail", "scanned_at"])
            overall = MaterialAttachment.ScanStatus.FAILED
            if getattr(settings, "FILE_SCAN_REQUIRED", False):
                raise
    return {"revision_id": revision_id, "status": overall}

def _build_docx(export):
    document = Document()
    document.add_heading(export.project.title, 0)
    document.add_paragraph(f"项目版本：{export.project_version}")
    document.add_paragraph(f"导出时间：{timezone.localtime().strftime('%Y-%m-%d %H:%M')}")
    manifest = []
    materials = export.project.materials.filter(status="approved").order_by("report_order", "id")
    for material in materials:
        revision = material.revisions.filter(status="approved").order_by("-created_at", "-id").first()
        if not revision:
            continue
        document.add_heading(material.report_section or material.title, level=1)
        document.add_heading(material.title, level=2)
        document.add_paragraph(revision.content)
        manifest.append({"material_id": material.id, "revision_id": revision.id, "title": material.title})
    output = BytesIO(); document.save(output)
    return output.getvalue(), manifest


@shared_task(bind=True, autoretry_for=(Exception,), retry_backoff=True, max_retries=3)
def generate_report_export(self, export_id):
    from .models import ReportExport

    export = ReportExport.objects.select_related("project").get(pk=export_id)
    export.status = ReportExport.Status.PROCESSING
    export.error_message = ""
    export.project_version = timezone.now().strftime("%Y%m%d%H%M%S")
    export.save(update_fields=["status", "error_message", "project_version"])
    try:
        docx_bytes, manifest = _build_docx(export)
        output_bytes = docx_bytes
        extension = "docx"
        if export.format == ReportExport.Format.PDF:
            converter = settings.DOCUMENT_CONVERTER_URL.rstrip("/")
            response = requests.post(
                f"{converter}/forms/libreoffice/convert",
                files={"files": ("report.docx", docx_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
                timeout=90,
            )
            response.raise_for_status()
            output_bytes, extension = response.content, "pdf"
        export.file.save(f"project-{export.project_id}-{export.project_version}.{extension}", ContentFile(output_bytes), save=False)
        export.material_manifest = manifest
        export.status = ReportExport.Status.COMPLETED
        export.completed_at = timezone.now()
        export.save(update_fields=["file", "material_manifest", "status", "completed_at"])
        return {"export_id": export.id, "status": export.status}
    except Exception as exc:
        export.status = ReportExport.Status.FAILED
        export.error_message = str(exc)[:2000]
        export.save(update_fields=["status", "error_message"])
        raise


@shared_task(bind=True, autoretry_for=(Exception,), retry_backoff=True, max_retries=2)
def generate_ai_response(self, record_id):
    from .models import AIGenerationLog, ProjectTask, Material, MaterialAttachment, MaterialRevision

    record = AIGenerationLog.objects.select_related("project", "actor", "task", "material").get(pk=record_id)
    conversation_message = getattr(record, "conversation_message", None)
    record.status = AIGenerationLog.Status.PROCESSING
    record.error_message = ""
    record.save(update_fields=["status", "error_message"])
    if conversation_message:
        conversation_message.status = AIConversationMessage.Status.STREAMING
        conversation_message.save(update_fields=["status", "updated_at"])
        publish_conversation_event(conversation_message.id, "message.started", {"message_id": conversation_message.id})
    try:
        project = record.project
        if project is None:
            if record.workspace_mode != "opening":
                raise ValueError("研究或答辩 AI 记录必须绑定当前项目。")
            api_key = getattr(settings, "OPENAI_API_KEY", "")
            if api_key:
                client_kwargs = {"api_key": api_key}
                base_url = getattr(settings, "OPENAI_BASE_URL", "")
                if base_url:
                    client_kwargs["base_url"] = base_url
                template = AgentTemplate.resolve(record.agent_key, record.actor.school, record.actor.role) if record.agent_key else None
                system = template.system_instruction if template else DEFAULT_AI_INSTRUCTION
                rendered_prompt = render_agent_prompt(template, record) if template else record.prompt
                if record.agent_key == "proposal-topic":
                    system = (
                        f"{system} 你正在执行研究问题助手工作流。"
                        "只输出严格 JSON，必须包含 project_title、project_type、project_plan、"
                        "candidates、recommended_index、missing_information；candidates 必须正好 3 个，并包含 question、scope、why、evidence_plan、limitations 和 scores。"
                    )
                response = OpenAI(**client_kwargs).responses.create(
                    model=settings.OPENAI_MODEL,
                    instructions=system,
                    input=f"Skill 指令：{rendered_prompt}\n用户开题想法：{record.prompt}",
                )
                output = response.output_text
                model_name = settings.OPENAI_MODEL
            else:
                output = _demo_ai_response(record, [], [], record.prompt)
                model_name = "演示模式（未接入真实模型）"
            artifact = _artifact_fields(record, output)
            record.output = output
            record.artifact_payload = artifact["artifact_payload"]
            record.verification_items = artifact["verification_items"]
            record.model_name = model_name
            record.status = AIGenerationLog.Status.COMPLETED
            record.referenced_sources = []
            record.completed_at = timezone.now()
            record.save(update_fields=["output", "artifact_payload", "verification_items", "model_name", "status", "referenced_sources", "completed_at"])
            if conversation_message:
                _finish_conversation_message(conversation_message, output, artifact)
            return {"record_id": record.id, "status": record.status, "mode": "opening"}
        context_parts = [f"项目题目：{project.title}", f"项目类型：{project.project_type}"]
        referenced = []
        scope = record.context_scope or {}
        if scope.get("project_basics", True):
            context_parts.extend([f"研究问题：{project.problem}", f"初步方案：{project.plan}"])
        if scope.get("approved_materials"):
            for material in project.materials.filter(status="approved"):
                revision = material.revisions.filter(status="approved").order_by("-created_at", "-id").first()
                if revision:
                    context_parts.append(f"已通过材料《{material.title}》：{revision.content[:6000]}")
                    referenced.append({"kind": "material", "id": material.id, "title": material.title, "project_id": project.id})

        # 当前步骤上下文：让学生/教师的问题"带着自己在哪一步、写了什么"一起进模型
        if scope.get("current_task") and record.task_id:
            try:
                t = record.task or ProjectTask.objects.get(pk=record.task_id)
                context_parts.append(f"当前步骤：{t.stage_name} · {t.title}\n步骤说明：{t.description}")
                referenced.append({"kind": "task", "id": t.id, "title": f"{t.stage_name} · {t.title}", "project_id": project.id})
                if t.evidence_requirements:
                    context_parts.append("该步骤的证据要求：" + "; ".join(str(e) for e in t.evidence_requirements))
            except ProjectTask.DoesNotExist:
                pass
        if scope.get("current_material_draft") and record.material_id:
            try:
                m = record.material or Material.objects.get(pk=record.material_id)
                rev = m.revisions.order_by("-created_at", "-id").first()
                if rev and rev.content:
                    context_parts.append(f"当前材料《{m.title}》你已写的内容：{rev.content[:6000]}")
                    referenced.append({"kind": "material", "id": m.id, "title": m.title, "project_id": project.id})
            except Material.DoesNotExist:
                pass
        if scope.get("current_guidance") and record.material_id:
            try:
                m = record.material or Material.objects.get(pk=record.material_id)
                guidance = m.effective_guidance
                if guidance:
                    context_parts.append(f"这份材料的写作要求（指引）：{guidance[:3000]}")
            except Material.DoesNotExist:
                pass

        # 学生自选“相关步骤”：让 AI 同时参考其它步骤的说明与证据要求
        related_task_ids = [int(x) for x in scope.get("related_tasks", []) if str(x).isdigit()]
        if related_task_ids:
            for t in ProjectTask.objects.filter(project=project, pk__in=related_task_ids):
                context_parts.append(f"相关步骤《{t.stage_name} · {t.title}》说明：{t.description}")
                if t.evidence_requirements:
                    context_parts.append("该步骤的证据要求：" + "; ".join(str(e) for e in t.evidence_requirements))
                referenced.append({"kind": "task", "id": t.id, "title": f"{t.stage_name} · {t.title}", "project_id": project.id})

        # 学生自选“相关材料”：参考其它材料的最新草稿与指引
        selected_material_ids = [int(x) for x in scope.get("selected_materials", []) if str(x).isdigit()]
        if selected_material_ids:
            for m in Material.objects.filter(project=project, pk__in=selected_material_ids):
                rev = m.revisions.order_by("-created_at", "-id").first()
                if rev and rev.content:
                    context_parts.append(f"相关材料《{m.title}》最新内容：{rev.content[:6000]}")
                guidance = m.effective_guidance
                if guidance:
                    context_parts.append(f"材料《{m.title}》的写作要求：{guidance[:3000]}")
                referenced.append({"kind": "material", "id": m.id, "title": m.title, "project_id": project.id})

        # 附件正文：当前材料 + 所选材料 + 相关步骤材料的抽取文本（PDF/Word/图片OCR）
        material_ids_for_files = set()
        if record.material_id:
            material_ids_for_files.add(record.material_id)
        material_ids_for_files.update(selected_material_ids)
        for t in ProjectTask.objects.filter(project=project, pk__in=related_task_ids):
            if t.material_id:
                material_ids_for_files.add(t.material_id)
        if material_ids_for_files:
            for att in MaterialAttachment.objects.filter(
                revision__material_id__in=material_ids_for_files
            ).exclude(extract_status=MaterialAttachment.ExtractStatus.PENDING).exclude(
                extract_status=MaterialAttachment.ExtractStatus.FAILED
            ).select_related("revision__material"):
                if att.extracted_text:
                    mat_title = att.revision.material.title if att.revision and att.revision.material else "附件"
                    context_parts.append(f"文件《{att.original_name}》（来自材料《{mat_title}》）内容：{att.extracted_text[:6000]}")
                    referenced.append({
                        "kind": "attachment",
                        "id": att.id,
                        "title": att.original_name,
                        "material_id": att.revision.material_id if att.revision else None,
                        "project_id": project.id,
                    })

        # 一致性体检：通读项目全部材料（含草稿），用于跨步骤矛盾检测
        if scope.get("consistency"):
            for m in project.materials.all().order_by("report_order", "id"):
                rev = m.revisions.order_by("-created_at", "-id").first()
                if rev and rev.content:
                    stage = m.task.stage_name if m.task else ""
                    context_parts.append(f"材料《{m.title}》（{stage}）最新内容：{rev.content[:4000]}")
                    referenced.append({"kind": "material", "id": m.id, "title": m.title, "project_id": project.id})

        # 仅在 Agent 契约明确要求时，带入同一项目中已完成的近期 AI 草稿。
        # 保持小而可审计的窗口，既方便迭代，又不会把其他项目或无限历史送入模型。
        if scope.get("ai_history"):
            history = (
                AIGenerationLog.objects.filter(project=project, status=AIGenerationLog.Status.COMPLETED)
                .exclude(pk=record.pk)
                .exclude(output="")
                .order_by("-completed_at", "-created_at", "-id")[:3]
            )
            for prior in history:
                title = prior.purpose or prior.agent_key or f"记录 #{prior.id}"
                context_parts.append(f"既有 AI 草稿《{title}》：{prior.output[:2000]}")
                referenced.append({"kind": "ai_history", "id": prior.id, "title": title, "project_id": project.id})

        # 教师审核意见是材料版本的一部分；只传当前项目、真实存在且有审核人的近期反馈。
        if scope.get("teacher_feedback"):
            feedback_revisions = (
                MaterialRevision.objects.filter(material__project=project, reviewer__isnull=False)
                .exclude(review_comment="")
                .select_related("material")
                .order_by("-created_at", "-id")[:8]
            )
            for revision in feedback_revisions:
                context_parts.append(
                    f"教师对材料《{revision.material.title}》的反馈：{revision.review_comment[:1500]}"
                )
                referenced.append({
                    "kind": "teacher_feedback", "id": revision.id, "title": revision.material.title,
                    "material_id": revision.material_id, "project_id": project.id,
                })

        tmpl = AgentTemplate.resolve(record.agent_key, record.project.school, record.actor.role) if record.agent_key else None
        system = tmpl.system_instruction if tmpl else DEFAULT_AI_INSTRUCTION
        rendered_prompt = render_agent_prompt(tmpl, record) if tmpl else record.prompt
        # Paper type is an execution-level contract, not a convenience token in
        # an editable template.  Always provide the validated value for each of
        # the six paper Agents even if an administrator later removes {paper_type}
        # from a template prompt.
        paper_type_context = (
            f"\n论文类型：{record.paper_type}\n"
            if record.agent_key in PAPER_AGENT_KEYS else ""
        )
        api_key = getattr(settings, "OPENAI_API_KEY", "")
        if not api_key:
            record.output = _demo_ai_response(record, context_parts, referenced, rendered_prompt)
            artifact = _artifact_fields(record, record.output)
            record.artifact_payload = artifact["artifact_payload"]
            record.verification_items = artifact["verification_items"]
            record.model_name = "演示模式（未接入真实模型）"
            record.status = AIGenerationLog.Status.COMPLETED
            record.referenced_sources = referenced
            record.completed_at = timezone.now()
            record.save(update_fields=["output", "artifact_payload", "verification_items", "model_name", "status", "referenced_sources", "completed_at"])
            if conversation_message:
                _finish_conversation_message(conversation_message, record.output, artifact)
            return {"record_id": record.id, "status": record.status, "mode": "demo"}
        client_kwargs = {"api_key": api_key}
        base_url = getattr(settings, "OPENAI_BASE_URL", "")
        if base_url:
            client_kwargs["base_url"] = base_url
        client = OpenAI(**client_kwargs)
        response = client.responses.create(
            model=settings.OPENAI_MODEL,
            instructions=system,
            input=(
                f"用途：{record.purpose}\n" + "\n".join(context_parts)
                + paper_type_context
                + f"\nSkill 指令：{rendered_prompt}\n用户补充：{record.prompt}"
            ),
        )
        record.output = response.output_text
        artifact = _artifact_fields(record, record.output)
        record.artifact_payload = artifact["artifact_payload"]
        record.verification_items = artifact["verification_items"]
        record.model_name = settings.OPENAI_MODEL
        record.status = AIGenerationLog.Status.COMPLETED
        record.referenced_sources = referenced
        record.completed_at = timezone.now()
        record.save(update_fields=["output", "artifact_payload", "verification_items", "model_name", "status", "referenced_sources", "completed_at"])
        if conversation_message:
            _finish_conversation_message(conversation_message, record.output, artifact)
        return {"record_id": record.id, "status": record.status}
    except Exception as exc:
        record.status = AIGenerationLog.Status.FAILED
        record.error_message = str(exc)[:2000]
        record.save(update_fields=["status", "error_message"])
        if conversation_message:
            conversation_message.status = AIConversationMessage.Status.FAILED
            conversation_message.error_message = record.error_message
            conversation_message.save(update_fields=["status", "error_message", "updated_at"])
            publish_conversation_event(conversation_message.id, "message.error", {"error": record.error_message})
        raise


def purge_trashed_project_records(retention_days: int = 30, dry_run: bool = False):
    """Delete expired trash records after removing private files and writing an audit summary."""
    from .models import AuditEvent, MaterialAttachment, Project, ReportExport, UploadPart

    cutoff = timezone.now() - timedelta(days=retention_days)
    projects = list(
        Project.all_objects.filter(deleted_at__isnull=False, trashed_at__lte=cutoff)
        .select_related("school", "leader", "primary_teacher")
        .order_by("id")
    )
    ids = [project.id for project in projects]
    if dry_run:
        return {"purged": len(ids), "ids": ids, "retention_days": retention_days, "dry_run": True}

    with transaction.atomic():
        for project in projects:
            # FileField.delete is best-effort and does not alter the immutable
            # audit summary when an old object store key is already missing.
            for attachment in MaterialAttachment.objects.filter(revision__material__project=project):
                if attachment.file:
                    attachment.file.delete(save=False)
            for export in ReportExport.objects.filter(project=project):
                if export.file:
                    export.file.delete(save=False)
            for part in UploadPart.objects.filter(session__revision__material__project=project):
                if part.file:
                    part.file.delete(save=False)
            AuditEvent.objects.create(
                school=project.school,
                actor=project.leader,
                action=AuditEvent.Action.PROJECT_PURGED,
                changes={
                    "project_id": project.id,
                    "title": project.title,
                    "school_id": project.school_id,
                    "purged_at": timezone.now().isoformat(),
                    "retention_days": retention_days,
                },
            )
            project.delete()
    return {"purged": len(ids), "ids": ids, "retention_days": retention_days, "dry_run": False}


@shared_task
def purge_trashed_projects(retention_days: int = 30, dry_run: bool = False):
    """Celery entry point shared with the management command."""
    return purge_trashed_project_records(retention_days=retention_days, dry_run=dry_run)
